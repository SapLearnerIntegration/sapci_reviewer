# New file: src/api/reporting_engine.py

import logging
from typing import Dict, List, Any, Optional
from datetime import datetime
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import json
import os

logger = logging.getLogger(__name__)

class ReportingEngine:
    """Generates comprehensive reports for SAP Integration Suite analysis"""

    def __init__(self, analysis_results: Dict[str, Any], tenant_data: Dict[str, Any]):
        """
        Initialize with analysis results and tenant data
        
        Args:
            analysis_results: Dict containing analysis results
            tenant_data: Dict containing tenant information
        """
        self.analysis_results = analysis_results
        self.tenant_data = tenant_data
        self.report_path = "reports"
        self.ensure_report_directory()

    def ensure_report_directory(self):
        """Ensure reports directory exists"""
        if not os.path.exists(self.report_path):
            os.makedirs(self.report_path)

    def generate_executive_summary(self) -> Dict[str, Any]:
        """
        Generate executive summary report
        
        Returns:
            Dict containing report details and file paths
        """
        try:
            # Create document
            doc = Document()
            self._set_document_styles(doc)

            # Add title
            title = doc.add_heading('SAP Integration Suite Analysis - Executive Summary', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add timestamp
            timestamp = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            timestamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Add tenant information
            self._add_tenant_overview(doc)

            # Add key findings
            self._add_key_findings(doc)

            # Add risk summary
            self._add_risk_summary(doc)

            # Add recommendations overview
            self._add_recommendations_overview(doc)

            # Add charts
            self._add_summary_charts(doc)

            # Save document
            timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
            report_filename = f"executive_summary_{timestamp_str}.docx"
            report_path = os.path.join(self.report_path, report_filename)
            doc.save(report_path)

            return {
                "success": True,
                "message": "Executive summary generated successfully",
                "report_path": report_path
            }

        except Exception as e:
            logger.error(f"Error generating executive summary: {str(e)}")
            return {
                "success": False,
                "message": f"Error generating executive summary: {str(e)}",
                "report_path": None
            }

    def generate_detailed_report(self) -> Dict[str, Any]:
        """
        Generate detailed technical report
        
        Returns:
            Dict containing report details and file paths
        """
        try:
            # Create document
            doc = Document()
            self._set_document_styles(doc)

            # Add title
            title = doc.add_heading('SAP Integration Suite Analysis - Detailed Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add timestamp
            timestamp = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            timestamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Add sections
            self._add_security_analysis(doc)
            self._add_error_handling_analysis(doc)
            self._add_performance_analysis(doc)
            self._add_compliance_analysis(doc)
            self._add_adapter_analysis(doc)
            self._add_detailed_recommendations(doc)

            # Save document
            timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
            report_filename = f"detailed_report_{timestamp_str}.docx"
            report_path = os.path.join(self.report_path, report_filename)
            doc.save(report_path)

            return {
                "success": True,
                "message": "Detailed report generated successfully",
                "report_path": report_path
            }

        except Exception as e:
            logger.error(f"Error generating detailed report: {str(e)}")
            return {
                "success": False,
                "message": f"Error generating detailed report: {str(e)}",
                "report_path": None
            }

    def generate_compliance_report(self) -> Dict[str, Any]:
        """
        Generate compliance-focused report
        
        Returns:
            Dict containing report details and file paths
        """
        try:
            # Create document
            doc = Document()
            self._set_document_styles(doc)

            # Add title
            title = doc.add_heading('SAP Integration Suite Analysis - Compliance Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add timestamp
            timestamp = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            timestamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Add compliance sections
            self._add_compliance_overview(doc)
            self._add_compliance_details(doc)
            self._add_compliance_recommendations(doc)
            self._add_compliance_charts(doc)

            # Save document
            timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
            report_filename = f"compliance_report_{timestamp_str}.docx"
            report_path = os.path.join(self.report_path, report_filename)
            doc.save(report_path)

            return {
                "success": True,
                "message": "Compliance report generated successfully",
                "report_path": report_path
            }

        except Exception as e:
            logger.error(f"Error generating compliance report: {str(e)}")
            return {
                "success": False,
                "message": f"Error generating compliance report: {str(e)}",
                "report_path": None
            }

    def _set_document_styles(self, doc: Document):
        """Set document styles"""
        # Heading styles
        styles = doc.styles
        
        for level in range(1, 4):
            style_name = f'Heading {level}'
            style = styles[style_name]
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(16 - (level * 2))
            font.bold = True
            if level == 1:
                font.color.rgb = RGBColor(31, 73, 125)

        # Normal text style
        style = styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

    def _add_tenant_overview(self, doc: Document):
        """Add tenant overview section"""
        doc.add_heading('Tenant Overview', 1)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Add headers
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Value'

        # Add data
        metrics = [
            ('Total Integration Flows', len(self.tenant_data.get('iflows', []))),
            ('Total Packages', len(self.tenant_data.get('packages', []))),
            ('Deployed Flows', sum(1 for iflow in self.tenant_data.get('iflows', []) if iflow.get('deployed', False))),
            ('Overall Health Score', f"{self.analysis_results.get('overall_health_score', 0)}%")
        ]

        for metric, value in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = str(value)

    def _add_key_findings(self, doc: Document):
        """Add key findings section"""
        doc.add_heading('Key Findings', 1)

        findings = self.analysis_results.get('key_findings', [])
        if not findings:
            doc.add_paragraph('No key findings to report.')
            return

        for finding in findings:
            p = doc.add_paragraph()
            p.add_run('• ').bold = True
            p.add_run(finding.get('description', ''))

    def _add_risk_summary(self, doc: Document):
        """Add risk summary section"""
        doc.add_heading('Risk Summary', 1)

        risks = {
            'High': self.analysis_results.get('high_risks', []),
            'Medium': self.analysis_results.get('medium_risks', []),
            'Low': self.analysis_results.get('low_risks', [])
        }

        for severity, risk_list in risks.items():
            if risk_list:
                doc.add_heading(f'{severity} Risks', 2)
                for risk in risk_list:
                    p = doc.add_paragraph()
                    p.add_run('• ').bold = True
                    p.add_run(risk.get('description', ''))

    def _add_recommendations_overview(self, doc: Document):
        """Add recommendations overview section"""
        doc.add_heading('Recommendations Overview', 1)

        recommendations = self.analysis_results.get('recommendations', {})
        categories = ['security', 'error_handling', 'performance', 'compliance']

        for category in categories:
            category_recommendations = recommendations.get(category, [])
            if category_recommendations:
                doc.add_heading(category.replace('_', ' ').title(), 2)
                for rec in category_recommendations[:3]:  # Top 3 recommendations
                    p = doc.add_paragraph()
                    p.add_run('• ').bold = True
                    p.add_run(rec.get('description', ''))

    def _add_summary_charts(self, doc: Document):
        """Add summary charts section"""
        doc.add_heading('Analysis Charts', 1)

        # Create charts
        self._create_security_chart()
        self._create_error_handling_chart()
        self._create_compliance_chart()

        # Add charts to document
        for chart_name in ['security_chart.png', 'error_handling_chart.png', 'compliance_chart.png']:
            chart_path = os.path.join(self.report_path, chart_name)
            if os.path.exists(chart_path):
                doc.add_picture(chart_path, width=Inches(6))
                doc.add_paragraph()  # Add spacing

    def _create_security_chart(self):
        """Create security analysis chart"""
        security_data = self.analysis_results.get('security_analysis', {})
        
        # Prepare data
        categories = ['High', 'Medium', 'Low']
        values = [
            security_data.get('high_vulnerabilities', 0),
            security_data.get('medium_vulnerabilities', 0),
            security_data.get('low_vulnerabilities', 0)
        ]

        # Create chart
        plt.figure(figsize=(10, 6))
        plt.bar(categories, values, color=['red', 'orange', 'green'])
        plt.title('Security Vulnerabilities by Severity')
        plt.xlabel('Severity')
        plt.ylabel('Number of Vulnerabilities')

        # Save chart
        plt.savefig(os.path.join(self.report_path, 'security_chart.png'))
        plt.close()

    def _create_error_handling_chart(self):
        """Create error handling analysis chart"""
        error_data = self.analysis_results.get('error_handling_analysis', {})
        
        # Prepare data
        labels = ['With Error Handling', 'Without Error Handling']
        sizes = [
            error_data.get('flows_with_error_handling', 0),
            error_data.get('flows_without_error_handling', 0)
        ]

        # Create chart
        plt.figure(figsize=(10, 6))
        plt.pie(sizes, labels=labels, autopct='%1.1f%%', colors=['green', 'red'])
        plt.title('Error Handling Implementation')

        # Save chart
        plt.savefig(os.path.join(self.report_path, 'error_handling_chart.png'))
        plt.close()

    def _create_compliance_chart(self):
        """Create compliance analysis chart"""
        compliance_data = self.analysis_results.get('compliance_analysis', {})
        
        # Prepare data
        categories = list(compliance_data.get('compliance_by_category', {}).keys())
        scores = [compliance_data.get('compliance_by_category', {}).get(cat, 0) for cat in categories]

        # Create chart
        plt.figure(figsize=(10, 6))
        plt.bar(categories, scores, color='blue')
        plt.title('Compliance Score by Category')
        plt.xlabel('Category')
        plt.ylabel('Compliance Score (%)')
        plt.xticks(rotation=45)

        # Save chart
        plt.savefig(os.path.join(self.report_path, 'compliance_chart.png'))
        plt.close()

    def _add_security_analysis(self, doc: Document):
        """Add security analysis section"""
        doc.add_heading('Security Analysis', 1)
        security_data = self.analysis_results.get('security_analysis', {})

        # Add overview
        doc.add_paragraph(f"Overall Security Score: {security_data.get('overall_score', 0)}%")

        # Add vulnerabilities table
        doc.add_heading('Security Vulnerabilities', 2)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(['Severity', 'Count', 'Description', 'Recommendation']):
            header_cells[i].text = header

        # Add vulnerabilities
        vulnerabilities = security_data.get('vulnerabilities', [])
        for vuln in vulnerabilities:
            row_cells = table.add_row().cells
            row_cells[0].text = vuln.get('severity', '')
            row_cells[1].text = str(vuln.get('count', 0))
            row_cells[2].text = vuln.get('description', '')
            row_cells[3].text = vuln.get('recommendation', '')

    def _add_error_handling_analysis(self, doc: Document):
        """Add error handling analysis section"""
        doc.add_heading('Error Handling Analysis', 1)
        error_data = self.analysis_results.get('error_handling_analysis', {})

        # Add overview
        doc.add_paragraph(f"Overall Error Handling Score: {error_data.get('overall_score', 0)}%")

        # Add findings
        doc.add_heading('Error Handling Findings', 2)
        findings = error_data.get('findings', [])
        for finding in findings:
            p = doc.add_paragraph()
            p.add_run('• ').bold = True
            p.add_run(finding.get('description', ''))

    def _add_performance_analysis(self, doc: Document):
        """Add performance analysis section"""
        doc.add_heading('Performance Analysis', 1)
        performance_data = self.analysis_results.get('performance_analysis', {})

        # Add overview
        doc.add_paragraph(f"Overall Performance Score: {performance_data.get('overall_score', 0)}%")

        # Add bottlenecks
        doc.add_heading('Performance Bottlenecks', 2)
        bottlenecks = performance_data.get('bottlenecks', [])
        for bottleneck in bottlenecks:
            p = doc.add_paragraph()
            p.add_run('• ').bold = True
            p.add_run(bottleneck.get('description', ''))

    def _add_compliance_analysis(self, doc: Document):
        """Add compliance analysis section"""
        doc.add_heading('Compliance Analysis', 1)
        compliance_data = self.analysis_results.get('compliance_analysis', {})

        # Add overview
        doc.add_paragraph(f"Overall Compliance Score: {compliance_data.get('overall_score', 0)}%")

        # Add compliance issues table
        doc.add_heading('Compliance Issues', 2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(['Category', 'Issue', 'Recommendation']):
            header_cells[i].text = header

        # Add issues
        issues = compliance_data.get('issues', [])
        for issue in issues:
            row_cells = table.add_row().cells
            row_cells[0].text = issue.get('category', '')
            row_cells[1].text = issue.get('description', '')
            row_cells[2].text = issue.get('recommendation', '')

    def _add_adapter_analysis(self, doc: Document):
        """Add adapter analysis section"""
        doc.add_heading('Adapter Analysis', 1)
        adapter_data = self.analysis_results.get('adapter_analysis', {})

        # Add overview
        doc.add_paragraph(f"Total Adapters: {adapter_data.get('total_adapters', 0)}")

        # Add adapter distribution
        doc.add_heading('Adapter Distribution', 2)
        distribution = adapter_data.get('distribution', {})
        for adapter_type, count in distribution.items():
            doc.add_paragraph(f"{adapter_type}: {count}")

    def _add_detailed_recommendations(self, doc: Document):
        """Add detailed recommendations section"""
        doc.add_heading('Detailed Recommendations', 1)
        recommendations = self.analysis_results.get('recommendations', {})

        for category, category_recommendations in recommendations.items():
            doc.add_heading(category.replace('_', ' ').title(), 2)
            for rec in category_recommendations:
                p = doc.add_paragraph()
                p.add_run(f"Priority: {rec.get('priority', 'Medium')}").bold = True
                p.add_run('\n')
                p.add_run(rec.get('description', ''))
                if rec.get('steps'):
                    p.add_run('\nImplementation Steps:')
                    for step in rec['steps']:
                        p = doc.add_paragraph()
                        p.add_run('• ').bold = True
                        p.add_run(step)